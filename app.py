import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import io
import zipfile
from typing import Dict, List, Tuple
import re

def copy_docx_template():
    """Create a new document with Roboto font style and logo header"""
    doc = Document()
    
    # Set default font to Roboto 9pt for the entire document
    try:
        style = doc.styles['Normal']
        style.font.name = 'Roboto'
        style.font.size = Pt(9)  # 9pt font size
    except:
        # Fallback if font setting fails
        pass
    
    # Add logo to header
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add logo image to header
        run = header_para.runs[0] if header_para.runs else header_para.add_run()
        run.add_picture('violet_logo.png', width=Inches(1.25))
    except Exception as e:
        # Fallback if logo fails to load
        print(f"Logo loading failed: {e}")
        pass
    
    return doc

def create_letter_from_template(mp_data):
    """Create a letter for an MP using the template content and replacing placeholders"""
    doc = copy_docx_template()
    
    # Replace placeholders with actual data first
    percentage = mp_data.get('percentage_electorate', '').strip()
    if percentage and percentage != 'nan':
        percentage_text = f"{percentage} of your constituents"
    else:
        percentage_text = "Many of your constituents"
    
    # Add date - bold
    date_para = doc.add_paragraph()
    date_run = date_para.add_run("4 September 2025")
    date_run.bold = True
    
    # Add full salutation
    doc.add_paragraph(mp_data['full_salutation'])
    
    # Add salutation - bold
    sal_para = doc.add_paragraph()
    sal_run = sal_para.add_run(mp_data['salutation'])
    sal_run.bold = True
    
    # Introduction paragraph
    intro_para = doc.add_paragraph()
    intro_para.add_run("I am writing to brief you on the ")
    intro_bold = intro_para.add_run("CARE Index 2025")
    intro_bold.bold = True
    intro_para.add_run(" - critical new research that directly impacts your constituents and presents immediate policy opportunities.")
    
    # The Challenge heading - bold
    challenge_para = doc.add_paragraph()
    challenge_run = challenge_para.add_run("The Challenge")
    challenge_run.bold = True
    
    # Challenge content
    challenge_content = doc.add_paragraph()
    challenge_content.add_run("The CARE Index 2025 (Caregiving and Ageing Readiness Evaluation) is Australia's first comprehensive assessment of our national preparedness for a rapidly ageing population. This landmark study commissioned by Violet reveals a critical finding: Australia scored just 23.1 out of 100 across five key pressure points - preparedness, planning, emotional toll, relationship strain, and competing demands. We are dangerously underprepared for what lies ahead.")
    
    # Why This Matters heading - bold
    matters_para = doc.add_paragraph()
    matters_run = matters_para.add_run("Why This Matters")
    matters_run.bold = True
    
    # Why This Matters content
    matters_content = doc.add_paragraph()
    matters_content.add_run("This result directly impacts your constituents and presents immediate policy opportunities.")
    
    # Bullet points
    bullet_points = [
        f"{percentage_text} are \"Sandwich Generation\" voters managing competing care responsibilities for elderly parents and their own children while attempting to remain in the workforce.",
        "88% feel unprepared and overwhelmed by current or anticipated caregiving demands.\nSubmissions from your constituents have been provided in this package.",
        "Almost 60% of Australians are already caring or will be within a decade, with two-thirds being women who face a 42% lifetime earnings gap while performing the equivalent of a 31.7 hour unpaid second job.",
        "$6+ billion annual economic impact from inappropriate care settings, preventable hospitalisations and low-value care due to inadequate planning, guidance and family support.",
        "Mental health crisis among family caregivers is creating workplace absence, relationship breakdown, and financial hardship, affecting hundreds of thousands of voters."
    ]
    
    for bullet in bullet_points:
        bullet_para = doc.add_paragraph()
        bullet_para.style = 'List Bullet'
        
        # Handle bold text in bullets
        if bullet.startswith(percentage_text):
            bold_run = bullet_para.add_run(percentage_text)
            bold_run.bold = True
            bullet_para.add_run(bullet[len(percentage_text):])
        elif bullet.startswith("88%"):
            bold_run = bullet_para.add_run("88% feel unprepared and overwhelmed")
            bold_run.bold = True
            bullet_para.add_run(bullet[len("88% feel unprepared and overwhelmed"):])
        elif bullet.startswith("Almost 60%"):
            bold_run = bullet_para.add_run("Almost 60% of Australians are already caring or will be within a decade,")
            bold_run.bold = True
            bullet_para.add_run(bullet[len("Almost 60% of Australians are already caring or will be within a decade,"):])
        elif bullet.startswith("$6+"):
            bold_run = bullet_para.add_run("$6+ billion annual economic impact")
            bold_run.bold = True
            bullet_para.add_run(bullet[len("$6+ billion annual economic impact"):])
        elif bullet.startswith("Mental health"):
            bold_run = bullet_para.add_run("Mental health crisis among family caregivers")
            bold_run.bold = True
            bullet_para.add_run(bullet[len("Mental health crisis among family caregivers"):])
        else:
            bullet_para.add_run(bullet)
    
    
    # Single empty line
    doc.add_paragraph()
    
    # Why Act Now heading - bold
    act_para = doc.add_paragraph()
    act_run = act_para.add_run("Why Act Now")
    act_run.bold = True
    
    # Why Act Now content
    act_content = doc.add_paragraph()
    act_content.add_run("Australia will experience a 400% increase in people reaching 85 years by 2031. We have proven solutions ready for immediate implementation through existing infrastructure like the Carer Gateway (over $1 billion invested) and Violet's technology-enabled platform already serving 30,000+ Australians.")
    
    # The Opportunity heading - bold
    opp_para = doc.add_paragraph()
    opp_run = opp_para.add_run("The Opportunity")
    opp_run.bold = True
    
    # The Opportunity content
    opp_content = doc.add_paragraph()
    opp_content.add_run("This research provides you with media-ready solutions for high-priority constituent concerns, evidence-based policy frameworks, and the first comprehensive national data on Australia's fastest-growing voter issue.")
    
    # Closing paragraph
    closing = doc.add_paragraph()
    closing.add_run("I would welcome the opportunity to brief you further on how the CARE Index findings can support your advocacy for constituents facing these challenges.")
    
    # Yours sincerely - bold
    yours_para = doc.add_paragraph()
    yours_run = yours_para.add_run("Yours sincerely,")
    yours_run.bold = True
    
    # Add empty lines for signature space
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sig_para = doc.add_paragraph()
    sig_run = sig_para.add_run("Kate Carnell AO")
    sig_run.bold = True
    
    chair_para = doc.add_paragraph()
    chair_para.add_run("Chair, Violet")
    
    # Add empty lines before attachment note
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Attachment note (in document body) - 8pt italic
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run("Attachment: Campaign submissions, CARE Index 2025 Executive Summary, Violet briefing note")
    footer_run.font.size = Pt(8)
    footer_run.italic = True
    
    # Add proper document footer
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Attachment: Campaign submissions, CARE Index 2025 Executive Summary, Violet briefing note"
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

def create_html_letter(mp_data):
    """Create an HTML version of the letter"""
    percentage = mp_data.get('percentage_electorate', '').strip()
    if percentage and percentage != 'nan':
        percentage_text = f"{percentage} of your constituents"
    else:
        percentage_text = "Many of your constituents"
    
    html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>MP Letter - {mp_data['first_name']} {mp_data['last_name']}</title>
    <style>
        body {{
            font-family: 'Roboto', Arial, sans-serif;
            font-size: 10pt;
            line-height: 1.4;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        .date {{
            font-weight: bold;
            margin-bottom: 20px;
        }}
        .salutation {{
            font-weight: bold;
            margin: 15px 0;
        }}
        .section-heading {{
            font-weight: bold;
            margin: 15px 0 10px 0;
        }}
        .intro-highlight {{
            font-weight: bold;
        }}
        ul {{
            margin: 15px 0;
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 10px;
        }}
        .bullet-bold {{
            font-weight: bold;
        }}
        .signature {{
            margin-top: 30px;
        }}
        .signature-name {{
            font-weight: bold;
            margin-top: 60px;
        }}
        .footer {{
            margin-top: 60px;
            font-style: italic;
        }}
        .header-logo {{
            float: right;
            width: 150px;
            height: auto;
            margin-bottom: 20px;
        }}
        .clear {{
            clear: both;
        }}
    </style>
</head>
<body>
    <div class="header-logo">
        <img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAR0AAABoCAYAAADM8iS9AAARs0lEQVR42u2debxWRRnHv6yXy2VxgRQEURRDQXErFHBBQQmFLMrcyCVSQzHTMi0zl9wKCtdwLbVwTS0NRdyyUDBc0CQwyg3FNERAglC5/fE89+P19i4zc86c97y+z/fzOR8u75k5Z86cM7/ZnnmmVWNjI4ZhGFnR2rLAMAwTHcMwTHQMwzBMdAzDMNExDMMw0TEMw0THMAwTHcMwDBMdwzBMdAzDMEx0DMOoDtrGuOjZh8yp9XzdAxgMbK3C/grwDDAL+Mg+uyB2BUYCfwfutOzIjvNu3y3/olPjDAWOUNHZQkVnCbAd0BG4y7LIi87AEGCUis7DJjrW0jE+Lhy/BMYVOLedHkcA84DjteVjFG/V7A0MA4YDXZqde8eyx0THgK7As8CWjgXqaa29n7SsA8233VVg9gH6lgi7zLLLRMeABxwFpzmPAhsDq2swvzoBn9fu0nDtihomOoYjBwAhI211wA+BM2osv4YAf7Rvr3axKfPknJAg7rgazK9/muCY6BjJGJAgbl+gW43l11vAz+yzMdExwmmfMP8712Ce/dw+GxMdI5x/J4i7AryjBvNsCWIoaZjoGAHclyDuPGBdjebbhfbpmOgY2XcVzqvhfHsMWdJgmOgYniwDDgmIdwXweI3n3U/s8zHRMcK4AxgP/Ncx/GRgkmVboq6pYaJT8/wa2Aq4FFhc4PxSYDtiHPddyy4A2lkW1B5mpJUubwCn6LEN0BtoBfwLeBFYb1lk319GdATaAKvspdcOL+lhGFmyg3b1jwWeAr5gomMYRgz2Bk4GvtTst1xau5voGEZ1szPwI2BsgXOr8phgEx3DqE7qkVnQidWWcBMdw6g+DgSuBnpWY+JtytwwqoupwL3VKjjW0nGjDnejP8OIxYbA75CdRqoaE53/pxviOnMfxCtgN2BH0lkN3hbYBdgf6AecCKzMsFW7mT7PBohLjTZ6bq2mYxlixLjCPgNAHMI35VkXPnZj8h/gPcT+6jXi219thSyZ6ZnhtzIGmQl7ArimGkTnQMRBVdIV1K2AD4DrImfwLsB+iL/eYdq6ac5JwJmB199CBazp6NHs3I8jik5nfZa9kBmOgS3uXYzViEX108Bc4CHE61+tiMwIYF+taPoj/pxLsRRxyv9nYAYwP+U0DdT30DGD598esfE5DOilv6XuCD+W6AwDvpfStf4bQXT68PE+SkMcCuNoD9HpgthM7KsiM7BE2BhuLQ4EvqUtqC4B8RuAQXocq7/NRZZ53EhOp2ETshcwAfgi/k7VeugxGnHX8QwyyJtG66Cfin/7iM++IXCois3uLc6tBGZWi+icgezzNCaFa50TSdGneYTfVj/GVUVaY4O1pbSn1pCV8AY4AVlEukOEaw/W4zzgWm2hfRrEZ2/EOf4+KV5zZxWd0/Xb/XXgdbqq2McSnFFaOY0p0Zp7Xo/UuxaxODGl69wSIW33AS94hG+nTe2WTAZeR/avOldbN1kLznDEGdi1kQSnZa14una3vl7FYtOg+fVoyoLTchzmZmQ30pCu0eOa32mW7wHA+cAi4H7tRpXqPj6F+LSuGtF5HdnxMgl/AV6NlD5fz3WbFanVNqtg4fkp8AgyJlWKG5FBwf7IGNPOiFFZaNO5m3Z573YY88gbwxDnYRPKhFuEWPoORfY064vMHJ3tWWGNAxYAn/WIc01KFci7QAfgSGRs7q/AWchiZBdejPECYs9enQYck7NWThP3AO97FJqGAr9N0heZNfXIBn97lgk3T/vqC1v8/ioy+PkLLVRTkZ1HfTlIP8wRVIcXwPHATQ7hzgIuKPD7y8iA8fnI1tCuXfQ+yADzIBWzUowGvpPS8+6GzK51D4y/MMZLiG0cuByYkiD+rRHTthZ40CN8qyI1weyMC05X7WeXE5zZwOccPpymcJcHpmdzTc92ORackxwFZ0wRwWnJ1cBOuE8G1OkYTanB/TYpf/M9EgjOWs8WXW5EB2Sg7oOAeLOQ6ciYPO0R9v0iv1+cYcFpA8wBti4TbqW2Pnw4GfhBYLo6aIHqlVPBOcJRVI/Hz5vhc8iYmk+FUaqim0J+tiSaT6TJgixEZ40Kjy8zMkibT591SZHf70OMxLLgYQoPaBcqPGsDrn+hYy1fiE7a9cgbO+E2g/QwYdPcTyDjPK4MprD5xcbAt1J+9hXIFP5C4BX9Tpcjxo3lDBqjbY2UlUXyZGSa1ed+o3WsISaLHcOto7RDrsv1+WJyPmJP4vJMSZroZyEmBWMD4vYBfqMtodzQGhlAdeHohO9mIrCph7hfD7zd7LdViAV8+xLds+vwG7ifi9hrNW8p1+lRr/eqR8YrG7Tr11FbZPOqXXQ+0oJwpEeckdqNWBwxXcuBxiLjNS0LcinL4WnaNamPlM7tVAxcuDqF+x2EbCK4UUDcw4EbtOVQaW5wfIY7SrRkXbkI8Y/tyiV8cpJlnUPr/gpP0WlToBz+R4/llawJsiJkoPKEyGlq1KMcz5Q5vwyxe4iFj4HZrSnly1cSxL+JyjMIOMox7GUp3O9G/CzMj8ZvkLdjARHJU/nOZaKeAv7hGecYh1ZIEhoc88BlcDHWIsn9dVzChb+lUGM38QiyqjmEnh4FPhbXOoZ7l3RmIFfgv1XyUdQgWSvhNM/wG2k/NxabO4ZzacXURUrjRR5hn0553klampXcvXQYYgbgwkzH1q4Lj3mGP9pEJz43Brzgb0VMT3/Hj3Jlhd7Prh6tHJAp3DR5K0FXaXM+OYiZJT6zpWkK9Z88ww+gshbtNSE67+A/FT7Co0Xii4tB25UVfD/Heob/e4ULcEsqsT5rE2TxrStpTlQ8j5iI+LAXNUYlBpquCIgzIVJaPu/Q37+3gu/Hd9p6WYQ0vBbQbWhiP7LfxdN3AHxJivdeg4yr+bB7rYlOJTwHPoBMx/rsyTMBPwMs1xqx3Hqjqyr4bnYKaHrH6gZeiriB8KUrslwjy+lzny7dh2W6Vx306KrHBoitSyfEcriz/t2p2W+9PdPbz0QnG64Bvu8RvgfiNiLNj3dfys+MTfW4XshgZKn7Dw6saWMwA/Eo2BDYmsxKdOrwW5awXlve9cikxYYqHg0tjrqIae5JjVEp0bnCU3RABpTT/HgPKnN+umd3xdfZUqPWtMUYFPBMsXz1rkOmzw8PiLt9ht/V9p7C2J70/D6F0rXWRKdSxkNLEQdKPoxJ8QXVOYjOOZEFfD2lF8LukLP3GSr4AzP8rrapwjLYMXJLykSnGSGzQmnNhoyl9ADnTPxngtLuXoVMpdZHfF+zAuP1zrBQbVmFZbAbNbYrSyUf9rfIIjefpfzHAz9L4d7fLHP+OxV+L52QQUtfukRM0+sqxP0C0tSddGeJitEnIM5ETVu9HnXaReug/++oFVTTubZ6tNF/W2s3rWlJTfsSlfkHyPqn1s3+foF4Y3EmOgW4AT/jv22QgcmnEtxzU0oPNs6kMt4Am9OZML8qG0VO15MBotOasAHoED4TEOcx/Ke5jSrtXoG4y0y7lVKOk8ucPzEH76VN4LvpEzldoda7HTLKtxBx28NkoLZEZ1HAh3x4wo+4lKhMx39Rap6IPWj7cs6/s5D79DMZqC3RAX8DvPaEL5Q7ssy4xyk5eS+hCxA/Gzldod7k1mWUbyH3GWwyUHuic0vAxzIx8F6lHGFNQdaG5YEmR0t5a+mE5k9WA6UhPn2HEncA3sih6KxBVp/7sD3+Rmd7lGgJrEF2Jc0LKwjzz9ONMKNCV95DLJN9WIusYcuCpYFlYLhJQW2JDoTZ7Pj6einl3+VUSlsHZ836BK2KERHTtRr/9V1vq1hlQejGjMd+Sst3Yx4TlRfRmY/sgujDUbivYO5P8QWL/8TfuVgWhE7j7hMxTesDWjoLMsyz1wLjjSVsut2oYtEB/+nzBmSrXBdKbatyZE7fzfzAeCOJu57H133sCxnmWRJ7m3M+heW7Qx4TlSfRuTmgi3OKQ5gtgS8XOTcDMXjLI88FxmsHfDFSmtrgbwvzXIZ5tpDw8aNv4r6FTLWQywHyPInOCvwdge9Oea+CF1VpX342YTujNhWgGDQEtKKy9KXzUcL73Z7j76FVQCtzE3K4I0TeEhSyFUip6fNewFeLnDuX7HbmDGEl7hvFtWQ34qy43gi/RaWzK5DHDyWIuwdwek6/h3UBlVB3cuiDOW+i8zj+BmhHlzh3bpHf36mSPvxtCeKeGyE9vt2P6RXIsztI5lfoEpLt+eXDL3HbHAAVnPcD7rGLiU55fHeo3AT4QpFauVj36RtV0ie/hXDDukMjjFH4rO1ap4Uqa5YDf0h4jduB8ZHT+WetMN/0iBMyXjXGRKc81wfEObXAb8VaMnMI30SuEk3qqQniX5dyenxqzUupnMuGC1K4xk3ATyOkbRfEi8FQZPLEx+7pzYD7HZa3cp5H0XkT/8HAEXxyi9bOFF/YeSjVxcUJugsHkO4WJ6McwzUCP65gns3VlkRSvqMCkcZsYE9kzHIest8V+G0XDWFbDNUnEOEt8HfDW5WiAzA5IE5zC+XvF3m2ywi3WnUpaL64zEasJNngZlozMlvivvTkVCq3QWHaXegBwD3aQp7k2cXsihge3oys0J/U7Nzb+HtjDBXSM/DbC2wbxE5sOhFsffLqJvEB7Ztv6BHn68D5WpC/XeD86iLdsLTwtTHyse6dAhxH2IzUZ4BbU2jhHecY7nmPLmHIota1juEWahcvrR1iB+vR1Fp5FnhFxaMpTR2AjZEZowHI1sadS3QBfSuqRxOkfyZi13ZZifsORMZBm8rPt2NUHnn2zToNONMjfB/ExedICvvknYjYccTCdyuRVp61yChkyUYIX9VCckmCFvFJjq29UR7XDbGc3tgj7CnAgcBWKb/rXSm/Z1oplhNmHvIeYoYwNPC+U7XifRCxFF+FuGPdSq/ZfPPJDxDPntSS6FzjKTpojV7IWHAB4Xtyu7A1/h7o2mrrzLW//TKyg8U9gWm8WGutEG+NVyJ+m8sxEr+V3iHGmSOQWbm3HMPvhazJytNQwuEJ4k5NIDpo+XDZMXdarC5yXsd00KbrbM84+wPbFvj9iMhpDXVx6ms5/DuSWRtfhb990kG4regfh/8EQMi4S1v8/Cm9kbCQps2VOnwQyp1kY3B5dqwL51l0CKyVW3Izcdf/tCV8r/XN8NsGt6kGSrK3+48QO5a+DmHHA3c7hBsD3OWZjsMIdyR/nGf4OcAQsvNgWIyHHLup5TghcjonEdEdSd5F5zbCrDCb90snRk7jwY5dj2J8LyDO9chsxLLAe44GFmutuyefdBHSDnGPcYdDl3QB4jTsvoA0nJYgzzbBfxr7SWT2bUGFvuUHtfuZBvcA90ZK5wPIDrzRyLvofIhY5YZyWkLRciHpWp3hhO1nPQvxhHhr4H1bqSD/Ubsg87VFuES7SQeXiT8FmaF5PuDe/Ulunv/dgDgvaZqvyvg7vjygRVuOsREEdC6FrftrSnRApj1DeFVfdkx2BHZK4TqhTe5l2k05gGR7gXVHtjEeRHlnVr/X506yIWEa09hDHbuIhThR4z8e+ftYrC2ykyNdf1fgiZSudS+yULg61SA6LwYq+vgM0nZpStc5k7AdPZuYgdiQjNamd9puKt9Fxtd21kI0P8G1eqU4JvHzBHGfQGa29ouQZ4t0XKSfinQs1qh4JrH+fl9FcWxWBbpa9lA+ncIGf8V4DvhTBulaivuMTVsdLyn0cTdoYXwvYXru16M3MrU8SmtD3xbBhyosc/R6j+HvprQYPfS6q1PK/6TM0qO3CtBoxKivt+d15mvL6W6SGfGF8EPgV8jM5jhk+YJLGbkFscX5d5aJbdXYmL7v5rMPmYORK7ZFDMB6IcZ1GyBrctoh1rSrEYO1t5A9y18ibHHhp4XWiHXu1sgMY3fEsrgDYkm+ViuItxEboEWEG27GYIi+897IDGE7bRUtQ0xRXsBjLO6823fLv+gYhmGUUnTDMAwTHcMwTHQMwzBMdAzDMNExDMMw0TEMw0THMAwTHcMwDBMdwzBMdAzDMEx0DMMw0TEMwzDRMQzDRMcwDBMdwzAMEx3DMKqd/wGbzBpjFnSrxAAAAABJRU5ErkJggg==" alt="Violet Logo" style="width: 150px; height: auto;">
    </div>
    <div class="clear"></div>
    
    <div class="date">4 September 2025</div>
    
    <div style="margin: 20px 0;">{mp_data['full_salutation']}</div>
    
    <div class="salutation">{mp_data['salutation']}</div>
    
    <p>I am writing to brief you on the <span class="intro-highlight">CARE Index 2025</span> - critical new research that directly impacts your constituents and presents immediate policy opportunities.</p>
    
    <div class="section-heading">The Challenge</div>
    <p>The CARE Index 2025 (Caregiving and Ageing Readiness Evaluation) is Australia's first comprehensive assessment of our national preparedness for a rapidly ageing population. This landmark study commissioned by Violet reveals a critical finding: Australia scored just 23.1 out of 100 across five key pressure points - preparedness, planning, emotional toll, relationship strain, and competing demands. We are dangerously underprepared for what lies ahead.</p>
    
    <div class="section-heading">Why This Matters</div>
    <p>This result directly impacts your constituents and presents immediate policy opportunities.</p>
    
    <ul>
        <li><span class="bullet-bold">{percentage_text}</span> are "Sandwich Generation" voters managing competing care responsibilities for elderly parents and their own children while attempting to remain in the workforce.</li>
        <li><span class="bullet-bold">88% feel unprepared and overwhelmed</span> by current or anticipated caregiving demands.<br><em>Submissions from your constituents have been provided in this package.</em></li>
        <li><span class="bullet-bold">Almost 60% of Australians are already caring or will be within a decade,</span> with two-thirds being women who face a 42% lifetime earnings gap while performing the equivalent of a 31.7 hour unpaid second job.</li>
        <li><span class="bullet-bold">$6+ billion annual economic impact</span> from inappropriate care settings, preventable hospitalisations and low-value care due to inadequate planning, guidance and family support.</li>
        <li><span class="bullet-bold">Mental health crisis among family caregivers</span> is creating workplace absence, relationship breakdown, and financial hardship, affecting hundreds of thousands of voters.</li>
    </ul>
    
    <div class="section-heading">Why Act Now</div>
    <p>Australia will experience a 400% increase in people reaching 85 years by 2031. We have proven solutions ready for immediate implementation through existing infrastructure like the Carer Gateway (over $1 billion invested) and Violet's technology-enabled platform already serving 30,000+ Australians.</p>
    
    <div class="section-heading">The Opportunity</div>
    <p>This research provides you with media-ready solutions for high-priority constituent concerns, evidence-based policy frameworks, and the first comprehensive national data on Australia's fastest-growing voter issue.</p>
    
    <p>I would welcome the opportunity to brief you further on how the CARE Index findings can support your advocacy for constituents facing these challenges.</p>
    
    <div class="signature">
        <div style="font-weight: bold;">Yours sincerely,</div>
        
        <div class="signature-name">Kate Carnell AO</div>
        <div>Chair, Violet</div>
    </div>
    
    <div class="footer">
        Attachment: Campaign submissions, CARE Index 2025 Executive Summary, Violet briefing note
    </div>
</body>
</html>
    """
    
    # Convert to BytesIO
    html_buffer = io.BytesIO()
    html_buffer.write(html_content.encode('utf-8'))
    html_buffer.seek(0)
    
    return html_buffer

def process_mp_csv(mps_df):
    """Process the MP CSV and generate individual letters for each MP"""
    results = {}
    
    for index, mp in mps_df.iterrows():
        # Create MP data dictionary
        mp_data = {
            'full_salutation': str(mp['full_salutation']).strip(),
            'salutation': str(mp['salutation']).strip(),
            'first_name': str(mp['First name']).strip(),
            'last_name': str(mp['Last name']).strip(),
            'electorate': str(mp['State/Electorate']).strip(),
            'percentage_electorate': str(mp['percentage_electorate']).strip() if pd.notna(mp['percentage_electorate']) and str(mp['percentage_electorate']).strip() != '' else ''
        }
        
        # Generate both DOCX and HTML letters
        doc_buffer = create_letter_from_template(mp_data)
        html_buffer = create_html_letter(mp_data)
        
        # Create a unique key for the MP
        mp_key = f"{mp_data['electorate']}_{mp_data['first_name']}_{mp_data['last_name']}"
        results[mp_key] = {
            'doc_buffer': doc_buffer,
            'html_buffer': html_buffer,
            'mp_data': mp_data
        }
    
    return results

def create_zip_file(mp_letters, format_type='both'):
    """Create a ZIP file containing letter files"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for mp_key, mp_info in mp_letters.items():
            mp_data = mp_info['mp_data']
            
            # Clean filename components
            electorate = re.sub(r'[<>:"/\\|?*]', '', mp_data['electorate'])
            first_name = re.sub(r'[<>:"/\\|?*]', '', mp_data['first_name'])
            last_name = re.sub(r'[<>:"/\\|?*]', '', mp_data['last_name'])
            base_filename = f"{electorate}, {first_name} {last_name} - MP Letter"
            
            if format_type in ['docx', 'both']:
                doc_buffer = mp_info['doc_buffer']
                docx_filename = f"{base_filename}.docx"
                zip_file.writestr(docx_filename, doc_buffer.getvalue())
            
            if format_type in ['html', 'both']:
                html_buffer = mp_info['html_buffer']
                html_filename = f"{base_filename}.html"
                zip_file.writestr(html_filename, html_buffer.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer

def main():
    st.title("MP Letter Generator")
    st.write("Upload the MP CSV file to generate individual letters for each MP using the CARE Index template.")
    
    # File upload section
    st.header("Upload MP Data")
    
    mps_file = st.file_uploader(
        "Upload MPs CSV File", 
        type=['csv'], 
        key="mps_file",
        help="CSV file containing MP information with required columns: full_salutation, salutation, First name, Last name, State/Electorate, percentage_electorate"
    )
    
    if mps_file is not None:
        try:
            # Load the CSV file
            mps_df = pd.read_csv(mps_file)
            
            # Validate required columns
            required_columns = ['full_salutation', 'salutation', 'First name', 'Last name', 'State/Electorate', 'percentage_electorate']
            missing_cols = [col for col in required_columns if col not in mps_df.columns]
            
            if missing_cols:
                st.error(f"CSV file is missing required columns: {', '.join(missing_cols)}")
                st.write("Expected columns:", required_columns)
                return
            
            # Display file information
            st.success("MP CSV file uploaded successfully!")
            st.info(f"**MPs loaded:** {len(mps_df)} records")
            
            # Show preview of the data
            st.write("Preview of MP data:")
            preview_df = mps_df[['First name', 'Last name', 'State/Electorate', 'percentage_electorate']].head(5)
            st.dataframe(preview_df)
            
            # Generate letters
            if st.button("Generate Letters", type="primary"):
                with st.spinner("Generating individual letters for each MP..."):
                    try:
                        mp_letters = process_mp_csv(mps_df)
                        
                        if not mp_letters:
                            st.warning("No letters were generated.")
                            return
                        
                        st.success(f"Generated {len(mp_letters)} letters!")
                        
                        # Show summary
                        st.header("Generated Letters Summary")
                        summary_data = []
                        
                        for mp_key, mp_info in mp_letters.items():
                            mp_data = mp_info['mp_data']
                            summary_data.append({
                                'MP Name': f"{mp_data['first_name']} {mp_data['last_name']}",
                                'Electorate': mp_data['electorate'],
                                'Percentage': mp_data['percentage_electorate'] if mp_data['percentage_electorate'] else 'N/A'
                            })
                        
                        summary_df = pd.DataFrame(summary_data)
                        st.dataframe(summary_df)
                        
                        # Download options
                        st.header("Download Letters")
                        
                        # Individual file downloads
                        with st.expander("Download Individual Letters"):
                            for mp_key, mp_info in mp_letters.items():
                                mp_data = mp_info['mp_data']
                                doc_buffer = mp_info['doc_buffer']
                                html_buffer = mp_info['html_buffer']
                                
                                # Create filenames
                                electorate = re.sub(r'[<>:"/\\|?*]', '', mp_data['electorate'])
                                first_name = re.sub(r'[<>:"/\\|?*]', '', mp_data['first_name'])
                                last_name = re.sub(r'[<>:"/\\|?*]', '', mp_data['last_name'])
                                base_name = f"{electorate}, {first_name} {last_name} - MP Letter"
                                
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.download_button(
                                        label=f"📄 DOCX: {mp_data['first_name']} {mp_data['last_name']} ({mp_data['electorate']})",
                                        data=doc_buffer.getvalue(),
                                        file_name=f"{base_name}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                                
                                with col2:
                                    st.download_button(
                                        label=f"🌐 HTML: {mp_data['first_name']} {mp_data['last_name']} ({mp_data['electorate']})",
                                        data=html_buffer.getvalue(),
                                        file_name=f"{base_name}.html",
                                        mime="text/html"
                                    )
                        
                        # Batch download as ZIP
                        st.subheader("Batch Download")
                        
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            zip_buffer_docx = create_zip_file(mp_letters, 'docx')
                            st.download_button(
                                label="📦 All DOCX Files",
                                data=zip_buffer_docx.getvalue(),
                                file_name="mp_care_index_letters_docx.zip",
                                mime="application/zip"
                            )
                        
                        with col2:
                            zip_buffer_html = create_zip_file(mp_letters, 'html')
                            st.download_button(
                                label="📦 All HTML Files",
                                data=zip_buffer_html.getvalue(),
                                file_name="mp_care_index_letters_html.zip",
                                mime="application/zip"
                            )
                        
                        with col3:
                            zip_buffer_both = create_zip_file(mp_letters, 'both')
                            st.download_button(
                                label="📦 All Files (Both Formats)",
                                data=zip_buffer_both.getvalue(),
                                file_name="mp_care_index_letters_all.zip",
                                mime="application/zip"
                            )
                        
                    except Exception as e:
                        st.error(f"Error generating letters: {str(e)}")
                        st.write("Please check your CSV file format and try again.")
        
        except Exception as e:
            st.error(f"Error loading CSV file: {str(e)}")
            st.write("Please ensure your file is a valid CSV format.")
    
    else:
        st.info("Please upload the MP CSV file to begin generating letters.")
        
        # Show expected file format
        with st.expander("Expected CSV File Format"):
            st.write("Your CSV file should contain the following columns:")
            st.code("""full_salutation,salutation,First name,Last name,State/Electorate,percentage_electorate
"Ms Jodie Belyea MP
Member for Dunkley
Member of Standing Committee on Health, Aged Care and Disability",Dear Ms Belyea,Jodie,Belyea,Dunkley,26.18%
"Senator Leah Blyth
Senator for South Australia",Dear Senator Blyth,Leah,Blyth,SA,""")
            
            st.write("**Column descriptions:**")
            st.write("- `full_salutation`: Complete formal title and position")
            st.write("- `salutation`: Simple greeting (e.g., 'Dear Ms Smith')")
            st.write("- `First name`: MP's first name")
            st.write("- `Last name`: MP's last name")
            st.write("- `State/Electorate`: State or electorate name")
            st.write("- `percentage_electorate`: Percentage for constituency data (can be empty for senators)")

if __name__ == "__main__":
    main()